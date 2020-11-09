"""Generates DSE Document based on dse Template information

Returns:
    [type] -- [description]
"""
import os
import xml.etree.ElementTree as ET
import const
import logger
import datetime
import toolbox

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from resources import Resources

class DocGenerator:
    """DocGenerator:
    """
   
    def __init__(self, checklist, name):
        self.checklistObject = checklist
        self.outputTemplate = {}
        self.outputDocument = None
        self.processed = False
        self.name = name

    def applyDocFormatting(self):
        style = self.outputDocument.styles['Normal']
        font = style.font
        font.name = "Arial"
        font.size = Pt(12)

        style = self.outputDocument.styles['Heading 1']
        font = style.font
        font.name = "Arial"
        font.size = Pt(16)

        style = self.outputDocument.styles['Heading 2']
        font = style.font
        font.name = "Arial"
        font.size = Pt(14)
        font.bold = True

        style = self.outputDocument.styles['Heading 3']
        font = style.font
        font.name = "Arial"
        font.size = Pt(12)
        font.bold = True


    def evaluateCondition(self, text):
        """evalutate conditions given in Mapping template for DSE Document
        
        Arguments:
            text {[type]} -- [description]
        """
        log = logger.getLogger()
        condition = False
        if text != "":
            try:
                condition = eval(text)
            except Exception (IndexError, OverflowError, SyntaxError, TypeError, NameError):                    
                log.warning("Error occured! Skipped evaluation! Condition to evaluate '" + text + "' will return False." )
        return condition

    def evaluateFormular(self, text):
        """evalutate formulars given in Mapping template for DSE Document
        
        Arguments:
            text {[type]} -- [description]
        """
        log = logger.getLogger()
        try:
            startpos = text.index("{")
        except(ValueError):
            return text
        if startpos > 0:
            endpos = text.index("}")+1
            if endpos > startpos:
                skript = text[startpos:endpos]
                cleanSkript = skript.replace("{","").replace("}","")
                try:
                    text = text.replace(skript, eval(cleanSkript)) 
                except (IndexError, OverflowError, SyntaxError, TypeError, NameError):                    
                    log.warning("Error occured! Skipped evaluation! String to evaluate '" + text + "' will be replaced by ''.")
                    text = text.replace(skript, "")    

                if text.count("{") > 0:
                    text = self.evaluateFormular(text)
            else:
                log.warning("Skipped processing! End position of symbol of text to evaluate is before start position (start:" + startpos +", end:" + endpos+")")
        return text

    def processText(self, text_elem):
        cond = ""
        text = text_elem.text
        cond = text_elem.attrib.get(const.DSEDOC_ATTRIB_COND)
        if cond != "" and cond != None:
            condition = self.evaluateCondition(cond)
            if condition is True:
                text = self.evaluateFormular(text)
            else:
                text = ""
        else:
            text = self.evaluateFormular(text)
        return text

    def processTable(self, table_elem):  #TODO
        text = table_elem.text
        #for elem in table_elem:
       #     if elem.tag == const.DSEDOC_TAG_ROW:

        return text


    def processParagraph(self, paragraph, chapter_title, is_first_paragraph):
        is_first = True
        title = paragraph.attrib.get(const.DSEDOC_ATTRIB_TITLE)
        text = paragraph.text
        if text != "":
            for elem in paragraph:
                if elem.tag == const.DSEDOC_TAG_TEXT:
                    text = self.processText(elem)
                elif elem.tag == const.DSEDOC_TAG_TABLE:
                    text = self.processTable(elem)            
                if text != "": 
                    # Print header only if there is content to print
                    if is_first_paragraph == True and is_first == True and chapter_title != "":
                        self.outputDocument.add_heading(chapter_title, level=1)
                    if is_first and title != "":    
                        self.outputDocument.add_heading(title, level=2)
                    paragraph = self.outputDocument.add_paragraph(text)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    is_first = False


    def processChapter(self, chapter):
        is_first_paragraph = True
        title = chapter.attrib.get(const.DSEDOC_ATTRIB_TITLE)
        for elem in chapter:
            if elem.tag == const.DSEDOC_TAG_PARAGRAPH:
                self.processParagraph(elem, title, is_first_paragraph)
                is_first_paragraph = False


    def parseTemplate(self, template):
        log = logger.getLoggerCtx("DSEGenerator.docGenerator.parseTemplate")
        filename = Resources.getOutputTemplate(template)
        try:
            tree = ET.parse(filename)   
        except(ET.ParseError):
            tree = None
            log.error("Error occured when parsing XML document '" + filename + "'! " + ET.ParseError.text)
        if tree is not None:
            root = tree.getroot()
            self.outputDocument = Document()
            core_properties = self.outputDocument.core_properties
            self.outputTemplate = template
            core_properties.comments = "Input Template:" + self.checklistObject.template + ",  Output Template: " + template
            self.applyDocFormatting()
            for elem in root:
                if elem.tag == const.DSEDOC_TAG_CHAPTER:
                    self.processChapter(elem)
            year = datetime.date.today().strftime("%Y")

            self.outputDocument.add_paragraph("@ " + year + " netvocat GmbH")
            self.processed = True 

    def saveDocument(self, versionnumber=1, path=None, now = datetime.datetime.now().strftime("%Y%m%d%H%M%S")):
        log = logger.getLoggerCtx("DSEGenerator.docGenerator.saveDocument")
        if path is None or len(path)==0:
            filename = Resources.getOutputPath() + "/" + now + "_" + self.outputTemplate[:-4] + ".docx"  
        else:
            filename = path  + "/" + now + "_" + self.outputTemplate[:-4] +".docx"  
        try:
            self.outputDocument.save(filename)
        except (PermissionError):
            log.warning("File '" + filename + "' could not be written! " + PermissionError.strerror)
            filename = Resources.getOutputPath() + "/" + now +  "_" + self.outputTemplate[:-4] +".docx"  
            self.saveDocument(filename)

        if os.path.isfile(filename):
            log.info("File '" + filename + "' has been written successfully!")
            return True
        
        log.warning("File '" + filename + "' has NOT been written! Please check error log!")
        return False


    