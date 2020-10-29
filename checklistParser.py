from docx import Document
from xmlObject import XMLObject
from resources import Resources
import xml.etree.ElementTree as ET
import const
import logging
import wx

def isWordTypeTable(child):
    return (child.attrib.get(const.CHECKLIST_ATTRIB_WORDTYPE) == const.CHECKLIST_ATTRIB_WORDTYPE_TABLE)

def isWordTypeCheckboxlist(child):
    return (child.attrib.get(const.CHECKLIST_ATTRIB_WORDTYPE) == const.CHECKLIST_ATTRIB_WORDTYPE_CHECKBOXLIST)

def isWordTypeYesNolist(child):
    return (child.attrib.get(const.CHECKLIST_ATTRIB_WORDTYPE) == const.CHECKLIST_ATTRIB_WORDTYPE_YESNOLIST)

def isWordTypeDynamicTable(child):
    return (child.attrib.get(const.CHECKLIST_ATTRIB_WORDTYPE) == const.CHECKLIST_ATTRIB_WORDTYPE_DYNAMICTABLE)

def isWordTypeTextInput(child):
    return (child.attrib.get(const.CHECKLIST_ATTRIB_WORDTYPE) == const.CHECKLIST_ATTRIB_WORDTYPE_TEXTINPUT)

def readVersion(checklistDocument):
    log = logging.getLogger("DSEGenerator.readVersion")
    try:
        version = checklistDocument.tables[0].table.cell(1, 0).text
        if version is None:
            version = "1.0"
        return version
    except (IndexError):
        log.error("list index out of range!")
        wx.MessageBox("Can't determine version of selected Document! Processing aborted. Please select another one!", caption="Error occured!")

def getTagContent(tag):
    text = ""
    text = tag.split(">")[1].split("<")[0]
    return text

def parseTextInput(xmlElement, checklistDocument):
    log = logging.getLogger("DSEGenerator.parseTextInput")
    element = ET.ElementTree(ET.fromstring(checklistDocument._element.xml))
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    e = element.find("w:sdt/w:sdtContent/w:p/w:r/w:t", namespaces)
    return e.text    

def parseYesNoList(xmlElement, checklistDocument):
    log = logging.getLogger("DSEGenerator.parseYesNoList")
    checkboxStatelist = []
    element = checklistDocument._element.xml
    elementliste = element.split("<")
    log.debug("doc: " + str(checklistDocument.text))
    for e in elementliste:
        if e.startswith("w14:checked "):
            e = e.split('"')[1]
            if e == "1" : 
                checkboxStatelist.append(True)
                break
            else:
                checkboxStatelist.append(False)
                break
    if checkboxStatelist is not None and len(checkboxStatelist)>0 and  checkboxStatelist[0] is True:
        return "True"
    else:
        return "False"

def parseCheckboxlist(xmlElement, checklistDocument):
    log = logging.getLogger("DSEGenerator.parseCheckboxlist")
    checkboxStatelist = []
    values = {}
    i = 0    
    element = checklistDocument._element.xml
    elementliste = element.split("<")
    for e in elementliste:
        if e.startswith("w14:checked "):
            e = e.split('"')[1]
            if e == "1" : 
                checkboxStatelist.append(True)
            else:
                checkboxStatelist.append(False)
    for e in checklistDocument.text.split("\n"):
        e=e.lstrip().rstrip()
        if i==0:
            i=i+1
            continue
        else:
            try:
                if len(e) > 0:
                    #Skip Remarks or Hints in Text
                    if e.startswith("Hinweis:"):
                        continue
                    else: 
                        log.debug("append key: " + e + "value: " + str(checkboxStatelist[i-1]) )
                        values[e] = checkboxStatelist[i-1]
            except (IndexError):
                log.error("list index out of range!")
        i=i+1
    return values

def parseDynamictable(xmlElement, checklistDocument):    #-- EXPERIMENTAL
    log = logging.getLogger("DSEGenerator.parseDynamictable")
    table = {}
    row = 0
    column = 0  
    row_data = {}  
    position = 0
    element = checklistDocument._element.xml
    elementliste = element.split("<w:tbl>")
    position = len(elementliste) - 1            
    table_elementlist = elementliste[position].split("<w:tr") #start with last table in given table
    for row_element in table_elementlist:
        row_data = {}
        for cell in row_element.split("<"):
            if cell.startswith("w:t>"):
                cell = getTagContent(cell)
                row_data[column] = cell
            column = column + 1
        if len(row_data) > 0:
            table[row] = row_data
            row = row + 1        
    return table


def parseTable(xmlElement, checklistDocument):
    zeile = 0
    pos = int(xmlElement.attrib.get(const.CHECKLIST_ATTRIB_TAB))
    table = checklistDocument.tables[pos-1]
    tableObject = {}    
    for child in xmlElement:
        if child.attrib.get(const.CHECKLIST_ATTRIB_COL) == None:
            spalte = 0
        else:
            spalte = int(child.attrib.get(const.CHECKLIST_ATTRIB_COL))-1       
        if isWordTypeTable(child):
            tableObject[child.tag] = parseTable(child, table.cell(zeile,spalte))
        elif isWordTypeCheckboxlist(child):
            tableObject[child.tag] = parseCheckboxlist(child, table.cell(zeile,spalte))
        elif isWordTypeYesNolist(child):
            tableObject[child.tag] = parseYesNoList(child, table.cell(zeile,spalte))
        elif isWordTypeDynamicTable(child):
            tableObject[child.tag] = parseDynamictable(child, table.cell(zeile,spalte))
        elif isWordTypeTextInput(child):
            tableObject[child.tag] = parseTextInput(child, table.cell(zeile, spalte))
        else:
            tableObject[child.tag] = table.cell(zeile,spalte).text      
        zeile = zeile + 1

    return tableObject


def parseChecklist(checklistFile):
    log = logging.getLogger("DSEGenerator.checklistParser")
    try:
        checklistDocument = Document(checklistFile)
    except (FileNotFoundError):
        log.error("File '" + checklistFile + "' not Found! " + FileNotFoundError.strerror)
        checklistDocument = None

    if checklistDocument is not None:
        version = readVersion(checklistDocument)
        checklistTemplate = Resources.getChecklisteTemplate(version)
        try:
            tree = ET.parse(checklistTemplate) 
        except Exception as e:
            log.error("Checklist Template '" + checklistTemplate + "' could not be read! " + str(e))
            tree = None
            checklistObject = None
        except (FileNotFoundError):
            wx.MessageBox("Error occured when opening Checklist Template! " + FileNotFoundError.strerror, caption="Error occured!")
            log.error("Checklist Template '" + checklistTemplate + "' could not be found! " + FileNotFoundError.strerror)

        if tree != None:
            root = tree.getroot()
            checklistObject = XMLObject()
            checklistObject.xmlVersion = root.attrib.get(const.DSEDOC_ATTRIB_VERSION)
            checklistObject.wordVersion = version
            if Resources.validVersions(version, checklistObject.xmlVersion):
                for elem in root:
                    if isWordTypeTable(elem):
                        checklistObject.addElement(elem.tag, parseTable(elem, checklistDocument))

                #if const.CHECKLIST_ATTRIB_TITLE in checklistObject.elementList:
                 #   checklistObject.wordVersion = checklistObject.elementList[const.CHECKLIST_ATTRIB_TITLE][const.CHECKLIST_ATTRIB_VERSION]
            else:
                log.warn("No valid versions! Processing skipped!")
    else:
        log.error("Processing of Template skipped! Please check Error log!")
        checklistObject = None
    log.debug("finished creation checklistObject: " + str(checklistObject))
    return checklistObject